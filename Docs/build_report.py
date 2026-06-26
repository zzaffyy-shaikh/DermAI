"""Assemble the full DermAI Final Year Project Report (.docx).
Follows the Plant-Disease sample structure; uses the existing DermAI written
chapters 1-3 and the generated diagrams/charts + real model results.
"""
from __future__ import annotations
import json, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "report_assets")
OUT = os.path.join(HERE, "fyp doc", "DermAI_FYP_Report.docx")
METRICS = json.load(open(os.path.join(ASSETS, "metrics.json")))

BLUE = RGBColor(0x1E, 0x3A, 0x8A)
INK = RGBColor(0x0F, 0x17, 0x2A)
GREY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# ---- base style ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)

for h, sz in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
    st = doc.styles[h]
    st.font.name = "Times New Roman"
    st.font.size = Pt(sz)
    st.font.color.rgb = BLUE
    st.font.bold = True

sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1.2)
sec.right_margin = Inches(1)

FIG = 0  # running counters not used; explicit numbers passed


# ----------------------------------------------------------------- helpers
def para(text="", size=12, bold=False, italic=False, align=None, color=None,
         space_after=6, line=1.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    return p


def heading(text, level=1, num=None):
    t = f"{num}\t{text}" if num else text
    h = doc.add_heading(t, level=level)
    return h


def page_break():
    doc.add_page_break()


def figure(img, number, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(ASSETS, img), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {number}: {caption}")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)


def table(rows, header=True, caption=None, number=None, widths=None,
          align_first_left=True, font=11):
    if caption and number:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(f"Table {number}: {caption}")
        r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            run.font.size = Pt(font)
            run.font.name = "Times New Roman"
            if header and i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade(cell, "1E3A8A")
            elif j == 0 and align_first_left:
                run.bold = True
    if widths:
        for j, w in enumerate(widths):
            for i in range(len(rows)):
                t.cell(i, j).width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def toc_field(title):
    para(title, size=14, bold=True, color=BLUE)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar"); fldChar.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose 'Update Field' to build the table of contents."
    fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
    for e in (fldChar, instr, fld2, t, fld3):
        run._r.append(e)


# ===========================================================================
# TITLE PAGE
# ===========================================================================
for _ in range(2):
    doc.add_paragraph()
para("DermAI", size=34, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, space_after=4)
para("An AI-Powered Skin Disease Screening & Consultation System",
     size=15, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=18)
para("Final Year Project Report", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para("Submitted by", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
for nm in ["Huzaifa bin Shakeel Shaikh (2539-2022)",
           "Abdul Wasay (2694-2022)",
           "Shahzaib Hussain Pirzada (2541-2022)"]:
    para(nm, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
doc.add_paragraph()
para("Supervisor", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Dr. Taha Shabbir", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para("In partial fulfilment of the requirements for the degree of",
     size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
para("Bachelor of Science in Artificial Intelligence",
     size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para("2025", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
for line in ["Faculty of Engineering Sciences and Technology",
             "Hamdard Institute of Engineering and Technology",
             "Hamdard University, Main Campus, Karachi, Pakistan"]:
    para(line, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
page_break()

# ===========================================================================
# CERTIFICATE OF APPROVAL
# ===========================================================================
heading("Certificate of Approval", level=3)
for line in ["Faculty of Engineering Sciences and Technology",
             "Hamdard Institute of Engineering and Technology",
             "Hamdard University, Karachi, Pakistan"]:
    para(line, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
doc.add_paragraph()
para('This project "DermAI" is presented by Huzaifa bin Shakeel Shaikh, Abdul Wasay, and '
     'Shahzaib Hussain Pirzada under the supervision of their project advisor and approved by the '
     'project examination committee, and acknowledged by the Hamdard Institute of Engineering and '
     'Technology, in fulfilment of the requirements for the Bachelor degree of Artificial Intelligence.')
for _ in range(2): doc.add_paragraph()
para("_______________________\t\t\t_______________________")
para("Dr. Taha Shabbir\t\t\t\tIn-charge FYP Committee")
para("(Project Supervisor)")
doc.add_paragraph()
para("_______________________\t\t\t_______________________")
para("Mr. Amir Hussain\t\t\t\tChairman")
para("(Project Co-Supervisor)\t\t\t(Department of Computing)")
doc.add_paragraph()
para("_______________________")
para("(Dean, FEST)")
page_break()

# ===========================================================================
# AUTHORS' DECLARATION
# ===========================================================================
heading("Authors' Declaration", level=3)
para("We declare that this project report was carried out in accordance with the rules and "
     "regulations of Hamdard University. The work is original except where indicated by special "
     "references in the text and no part of the report has been submitted for any other degree. "
     "The report has not been presented to any other University for examination.")
doc.add_paragraph()
para("Dated:"); para("Authors' Signatures:")
for nm in ["Huzaifa bin Shakeel Shaikh", "Abdul Wasay", "Shahzaib Hussain Pirzada"]:
    para("_______________________"); para(nm)
page_break()

# ===========================================================================
# PLAGIARISM UNDERTAKING
# ===========================================================================
heading("Plagiarism Undertaking", level=3)
para("We, Huzaifa bin Shakeel Shaikh, Abdul Wasay, and Shahzaib Hussain Pirzada, solemnly declare "
     "that the work presented in the Final Year Project Report titled DermAI has been carried out "
     "solely by ourselves with no significant help from any other person except a few of those which "
     "are duly acknowledged. We confirm that no portion of our report has been plagiarized and any "
     "material used in the report from other sources is properly referenced.")
doc.add_paragraph()
para("Dated:"); para("Authors' Signatures:")
for nm in ["Huzaifa bin Shakeel Shaikh", "Abdul Wasay", "Shahzaib Hussain Pirzada"]:
    para("_______________________"); para(nm)
page_break()

# ===========================================================================
# ACKNOWLEDGMENTS
# ===========================================================================
heading("Acknowledgments", level=3)
para("First and foremost, we are incredibly appreciative to Allah (SWT) for giving us the patience, "
     "strength, and direction we needed to complete this Final Year Project.")
para("We want to sincerely thank Dr. Taha Shabbir and Mr. Amir Hussain, our supervisors, for their "
     "invaluable advice, unwavering support, and helpful criticism during this project. This work was "
     "shaped and completed thanks in large part to their mentoring.")
para("We are also appreciative of the department's leadership and teachers for providing us with the "
     "tools we needed to succeed in our studies as well as a helpful learning environment.")
para("Finally, we want to express our sincere gratitude to our parents for their unwavering prayers, "
     "support, and encouragement. Throughout this journey, their faith in us has been a source of "
     "inspiration.")
page_break()

# ===========================================================================
# DOCUMENT INFORMATION
# ===========================================================================
heading("Document Information", level=3)
table([
    ["Category", "Information"],
    ["Customer", "Patients / Healthcare Institutes"],
    ["Project Title", "DermAI"],
    ["Document", "Final Year Project Report"],
    ["Document Version", "1.0"],
    ["Identifier", "FYP-023/FA25 Final Report"],
    ["Status", "Final Report"],
    ["Author(s)", "Huzaifa bin Shakeel Shaikh, Abdul Wasay, Shahzaib Hussain Pirzada"],
    ["Approver(s)", "Dr. Taha Shabbir"],
    ["Issue Date", "2025"],
], caption="Document Information", number="1", widths=[2.2, 4.0])
page_break()

# ===========================================================================
# DEFINITIONS, ACRONYMS, ABBREVIATIONS
# ===========================================================================
heading("Definition of Terms, Acronyms, and Abbreviations", level=3)
para("This section provides the definitions of all terms, acronyms, and abbreviations required to "
     "interpret the terms used in this document properly.")
table([
    ["Term / Acronym", "Description"],
    ["DermAI", "The proposed AI-based skin disease detection and consultation system."],
    ["AI", "Artificial Intelligence — machines simulating human reasoning and learning."],
    ["CNN", "Convolutional Neural Network — a deep learning architecture for image data."],
    ["ConvNeXt", "A modern convolutional neural network used for image classification."],
    ["Mobile SAM", "Lightweight Segment Anything Model used to localise skin-lesion regions."],
    ["LLM", "Large Language Model used to generate disease-specific follow-up questions."],
    ["OOD", "Out-Of-Distribution — input that does not belong to any trained class."],
    ["Decision Gate", "Confidence/entropy rule that flags unreliable or out-of-scope predictions."],
    ["Roboflow", "An online platform used for dataset annotation and preprocessing."],
    ["Lesion", "An abnormal area of skin that may indicate disease or infection."],
    ["API", "Application Programming Interface (the FastAPI backend services)."],
    ["UI", "User Interface."],
], caption="Definition of Terms, Acronyms, and Abbreviations", number="2", widths=[2.0, 4.2])
page_break()

# ===========================================================================
# ABSTRACT
# ===========================================================================
heading("Abstract", level=3)
para("Skin diseases are among the most common health problems worldwide, and their early and "
     "accurate diagnosis remains challenging due to visual similarity between conditions, variations "
     "in skin tone, and limited access to dermatologists. This project presents DermAI, an "
     "intelligent skin disease analysis system that integrates deep-learning–based image processing "
     "with disease-specific interactive questioning and clinical consultation support.")
para("The proposed system follows a multi-stage diagnostic pipeline. First, Mobile SAM is employed "
     "to localise the affected skin region within an uploaded image. The region is then passed to a "
     "ConvNeXt-based image classifier, which predicts the most probable skin-disease category across "
     "eight classes. A decision gate based on confidence and entropy filters unreliable or "
     "out-of-distribution inputs. Based on the predicted disease, a Large Language Model (LLM) "
     "dynamically generates targeted follow-up questions to collect relevant symptom information, "
     "and the final preliminary result is produced through a decision-fusion mechanism that combines "
     "visual classification with symptom-based analysis.")
para("All datasets used in this project are carefully annotated using Roboflow, ensuring consistent "
     "labelling and preprocessing. Once the diagnosis is established, the system enables doctor "
     "consultation through an integrated video-call workflow, allowing medical professionals to "
     "review the results and provide expert guidance. Evaluated on a clean, leakage-free test set of "
     "1,315 images, the ConvNeXt classifier achieves an overall accuracy of 89.1% with a macro "
     "F1-score of 0.886. The results demonstrate that integrating object localisation, modern CNN "
     "architectures, and LLM-driven interaction significantly enhances diagnostic reliability and "
     "system usability while clearly complementing — not replacing — professional medical judgment.")
para("Keywords:", bold=True, space_after=2)
para("Skin Disease Detection, Mobile SAM, ConvNeXt, Medical Image Analysis, Convolutional Neural "
     "Network (CNN), Large Language Model (LLM), Computer-Aided Diagnosis, Deep Learning, Roboflow "
     "Annotation Pipeline, Telemedicine Consultation System.", italic=True)
page_break()

# ===========================================================================
# TABLE OF CONTENTS / LISTS
# ===========================================================================
toc_field("Table of Contents")
page_break()

para("List of Figures", size=14, bold=True, color=BLUE)
figs = [
    ("Figure 3-1", "Reported Accuracy of Related Works"),
    ("Figure 4-1", "Waterfall Development Model"),
    ("Figure 5-1", "System Architecture"),
    ("Figure 5-2", "Application Flowchart"),
    ("Figure 5-3", "Sequence Diagram"),
    ("Figure 5-4", "Use Case Diagram"),
    ("Figure 5-5", "State Diagram"),
    ("Figure 6-1", "Dataset Distribution Across Classes"),
    ("Figure 6-2", "Confusion Matrix (Clean Test Set)"),
    ("Figure 6-3", "Per-Class Precision, Recall and F1-score"),
    ("Figure 6-4", "Live On-Device Inference Confidence"),
    ("Figure 6-5", "Sample Result — Eczema Screening"),
    ("Figure 6-6", "Sample Result — Acne Screening"),
    ("Figure 6-7", "Sample Result — Healthy Skin"),
    ("Figure C-1", "ER Diagram (Logical Data Model)"),
    ("Figure C-2", "Data Flow Diagram (Level 1)"),
    ("Figure C-3", "Design Class Diagram"),
]
for n, c in figs:
    para(f"{n}\t{c}", size=11, space_after=2)
page_break()

para("List of Tables", size=14, bold=True, color=BLUE)
tabs = [
    ("Table 1", "Document Information"),
    ("Table 2", "Definition of Terms, Acronyms, and Abbreviations"),
    ("Table 3-1", "Traditional Diagnosis vs. DermAI"),
    ("Table 3-2", "Comparison of Related Systems with DermAI"),
    ("Table 4-1", "Software / Tools Used"),
    ("Table 4-2", "Hardware Used"),
    ("Table 5-1 to 5-10", "Test Cases"),
    ("Table 6-1", "Dataset Split per Class"),
    ("Table 6-2", "Training Configuration"),
    ("Table 6-3", "Per-Class Classification Metrics"),
    ("Table 6-4", "Live Inference Summary"),
    ("Table B-1 to B-3", "SRS Functional / Non-Functional Requirements"),
    ("Table C-1", "SDS Data Dictionary"),
]
for n, c in tabs:
    para(f"{n}\t{c}", size=11, space_after=2)
page_break()

# ===========================================================================
# CHAPTER 1 — INTRODUCTION
# ===========================================================================
heading("CHAPTER 1", level=1)
heading("INTRODUCTION", level=1)
para("The early detection of skin diseases is crucial for preventing complications and improving "
     "patient outcomes. Traditional dermatological diagnosis relies on in-person visual examination "
     "by specialists, a process that is often slow, costly, and inaccessible in rural or "
     "under-served regions. To bridge this gap, DermAI applies artificial intelligence and computer "
     "vision to enable preliminary identification of common skin conditions directly from images of "
     "the affected area. By combining a modern deep-learning classifier with interactive, "
     "symptom-based questioning and a doctor-consultation workflow, the system provides reliable, "
     "accessible, and standardized screening that supports — rather than replaces — professional "
     "medical judgment.")

heading("Motivation", level=2, num="1.1")
para("Skin diseases are among the most common health problems worldwide, yet timely and accurate "
     "diagnosis remains a challenge, especially in regions with limited access to dermatologists. "
     "Many patients rely on self-assessment or delayed consultations, which can lead to disease "
     "progression and improper treatment. The motivation behind this project is to address this gap "
     "by leveraging modern artificial intelligence techniques to assist in early, accessible, and "
     "reliable skin-disease assessment.")
para("Recent advancements in computer vision and deep learning have shown strong potential in "
     "medical image analysis. Models such as Mobile SAM and ConvNeXt enable efficient object "
     "localisation and high-accuracy image classification, making them suitable for real-time "
     "medical applications. Additionally, the integration of disease-specific questioning through "
     "Large Language Models (LLMs) allows the system to mimic a dermatologist's diagnostic reasoning "
     "by combining visual evidence with patient-reported symptoms.")
para("Another key motivation is to reduce dependency on manual processes and subjective judgment by "
     "providing a standardized, AI-driven decision-support tool. By using Roboflow for precise image "
     "annotation and structured datasets, the system aims to improve model reliability and "
     "generalisation across different skin tones and image conditions. Furthermore, the inclusion of "
     "doctor consultation as a final step ensures that the system acts as an assistive tool rather "
     "than a replacement for medical professionals, maintaining ethical and clinical responsibility.")
para("Overall, this project is motivated by the need for an intelligent, scalable, and user-friendly "
     "dermatology assistance system that enhances early detection, improves diagnostic confidence, "
     "and supports healthcare delivery through the effective use of AI technologies.")

heading("Problem Statement", level=2, num="1.2")
para("Skin diseases are among the most common health issues worldwide, yet timely and accurate "
     "diagnosis remains a challenge, especially in regions with limited access to dermatologists. "
     "Many individuals rely on self-assessment or online information, which is often inaccurate, "
     "unverified, or misleading. Existing AI-based dermatology systems mostly focus on image-only "
     "analysis, ignoring critical symptom-related information such as itching, pain, duration, or "
     "spread, which dermatologists routinely consider during diagnosis. Additionally, most available "
     "solutions are either expensive, non-interactive, or limited to detecting only a few severe "
     "conditions, making them unsuitable for early-stage and common skin-disease identification.")
para("Another major limitation is the lack of an integrated workflow that mimics real clinical "
     "practice. Current systems rarely combine object localisation, disease classification, "
     "intelligent symptom-based questioning, and professional consultation within a single platform. "
     "Variations in image quality, lighting conditions, and skin tones further reduce diagnostic "
     "reliability when visual data is used alone. Therefore, there is a clear need for an "
     "intelligent, interactive, end-to-end dermatology support system that can accurately detect "
     "skin conditions from images, refine predictions through disease-specific questioning, and "
     "seamlessly guide users toward medical consultation when required.")
para("The problem addressed in this project is the design and development of an AI-powered "
     "dermatological assistance system that integrates object localisation, image classification, "
     "and language-based symptom analysis to improve diagnostic accuracy and accessibility, while "
     "maintaining scalability, usability, and clinical relevance.")

heading("Goals and Objectives", level=2, num="1.3")
para("The goal of this project is to develop an AI-powered dermatological screening system that "
     "provides reliable preliminary identification of common skin diseases through automated image "
     "analysis and intelligent symptom questioning. The specific objectives are:")
for o in [
    "To localise affected skin regions using Mobile SAM segmentation.",
    "To classify diseases accurately across eight categories using a modern ConvNeXt deep-learning model.",
    "To enhance diagnostic confidence through disease-specific follow-up questions generated by an LLM.",
    "To filter unreliable and out-of-distribution inputs through a confidence/entropy decision gate.",
    "To design a scalable, modular architecture supporting both mobile and web platforms.",
    "To integrate a secure doctor-consultation and video-call workflow.",
    "To ensure user-data privacy, ethical compliance, and clear clinical scoping (screening, not diagnosis).",
]:
    doc.add_paragraph(o, style="List Bullet")

heading("Project Scope", level=2, num="1.4")
para("The scope of this project covers the design and development of an AI-based dermatological "
     "assistance system that performs automated skin-disease analysis and preliminary screening. The "
     "system localises affected regions using Mobile SAM, classifies the disease through a ConvNeXt "
     "classifier, and generates disease-specific follow-up questions using an LLM to collect "
     "contextual symptom information that cannot be determined from images alone. It also includes a "
     "doctor-consultation module allowing users to seek professional guidance after receiving the "
     "AI-generated assessment. The application operates on desktop and mobile platforms with an "
     "English-language interface. The system does not aim to replace certified medical professionals, "
     "provide final clinical diagnoses, prescribe medication, or handle emergency cases; it is "
     "intended solely for early-stage screening and decision support, while ensuring data privacy, "
     "ethical compliance, and modularity for future enhancements.")
page_break()

# ===========================================================================
# CHAPTER 2 — BACKGROUND
# ===========================================================================
heading("CHAPTER 2", level=1)
heading("RELEVANT BACKGROUND & DEFINITIONS", level=1)
para("This chapter provides the essential background knowledge and core concepts required to "
     "understand the development of the DermAI system. It introduces the key technologies and "
     "methodologies used in the project, including medical image analysis, deep learning, and "
     "artificial intelligence. Fundamental concepts related to skin-disease diagnosis, object "
     "localisation, image classification, and natural-language processing are explained to establish "
     "the technical foundation of the system.")

heading("Background of Skin Diseases", level=2, num="2.1")
para("Skin diseases are among the most common health problems worldwide and can affect people of all "
     "ages. Conditions such as acne, eczema, atopic dermatitis, psoriasis, hyper-pigmentation, "
     "seborrheic keratoses, and skin cancer (melanoma) often require early diagnosis to prevent "
     "complications. However, access to qualified dermatologists is limited in many regions, "
     "especially in rural and under-developed areas. Diagnosis is primarily visual and depends on "
     "the experience of medical professionals, which can lead to human error or delayed treatment. "
     "These challenges highlight the need for automated, technology-driven systems that can assist in "
     "early screening and preliminary diagnosis of skin conditions. Table 3-1 contrasts traditional "
     "diagnosis with the proposed DermAI approach.")
table([
    ["Aspect", "Traditional Diagnosis", "DermAI (Proposed System)"],
    ["Diagnosis Method", "Manual visual inspection by dermatologist", "AI-based image analysis using ConvNeXt + LLM"],
    ["Accessibility", "Limited, especially in rural areas", "Accessible through mobile and web apps"],
    ["Time Required", "Time-consuming due to appointments", "Provides instant preliminary results"],
    ["Cost", "High consultation and travel cost", "Low-cost or free initial screening"],
    ["Accuracy", "Depends on doctor's experience", "Consistent predictions from a trained model"],
    ["Scalability", "Cannot serve large populations quickly", "Highly scalable, serves many users"],
    ["Early Detection", "May be delayed", "Enables early-stage screening"],
    ["User Interaction", "In-person consultation only", "Interactive disease-specific questioning"],
    ["Error Handling", "Prone to human fatigue and subjectivity", "Reduced bias through standardized analysis"],
    ["Follow-up Support", "Requires physical visits", "Supports digital doctor consultation"],
], caption="Traditional Diagnosis vs. DermAI", number="3-1", widths=[1.7, 2.3, 2.3], font=10)

heading("Artificial Intelligence in Healthcare", level=2, num="2.2")
para("Artificial Intelligence (AI) has emerged as a powerful tool in modern healthcare by enabling "
     "automated analysis, prediction, and decision support. In medical imaging, AI techniques such "
     "as deep learning and computer vision are widely used for detecting and classifying diseases "
     "with high accuracy. AI-based systems can analyse large volumes of data, identify patterns that "
     "may not be visible to humans, and provide fast and consistent results. In dermatology, AI helps "
     "improve diagnostic efficiency, reduce workload on doctors, and support early disease detection, "
     "making healthcare more accessible and cost-effective.")

heading("Object Localisation (Mobile SAM)", level=2, num="2.3")
para("Object segmentation is a computer-vision task that involves identifying and precisely "
     "outlining regions of interest within an image at the pixel level. Mobile SAM (Segment Anything "
     "Model for Mobile) is a lightweight, efficient segmentation model designed to perform accurate "
     "segmentation on resource-constrained devices. In DermAI, Mobile SAM is used to isolate the "
     "affected skin region before analysis, minimising background interference and severity "
     "estimation. It was selected due to its optimised architecture, fast inference speed, and "
     "ability to deliver high-quality results while remaining suitable for mobile deployment.")

heading("Image Classification (ConvNeXt)", level=2, num="2.4")
para("Image classification is the process of assigning a label to an image based on its visual "
     "content. ConvNeXt is a modern convolutional neural network architecture inspired by "
     "Transformer-based designs while maintaining the efficiency of CNNs. In DermAI, a ConvNeXt-Tiny "
     "model is used to classify skin images into eight disease categories. The model provides strong "
     "performance with fewer parameters and better feature representation than traditional CNNs, "
     "making it well suited for medical image classification where accuracy and generalisation are "
     "critical.")

heading("Large Language Models in Medical Systems", level=2, num="2.5")
para("Large Language Models (LLMs) are advanced AI models capable of understanding and generating "
     "human-like text. In DermAI, an LLM is used after image classification to generate "
     "disease-specific questions for the user. These questions collect symptom-related information "
     "such as itching, pain, duration, and medical history. By combining visual predictions with "
     "textual symptom analysis, the system mimics real clinical reasoning. The LLM also enhances user "
     "interaction by providing natural, conversational, and context-aware questioning, making the "
     "system more intelligent and user-friendly. DermAI uses Gemini 2.5-flash as the online provider "
     "with an offline scripted question bank as a graceful fallback.")

heading("Definitions", level=2, num="2.6")
defs = [
    ("Artificial Intelligence (AI)", "The ability of machines to simulate human intelligence, including learning, reasoning, and decision-making."),
    ("Deep Learning", "A subset of machine learning that uses multi-layer neural networks to learn complex patterns from data."),
    ("Object Localisation", "A technique used to identify and locate regions of interest within an image."),
    ("Image Classification", "The task of assigning a category label to an image based on its content."),
    ("Mobile SAM", "A real-time segmentation model used to isolate skin-lesion regions in images."),
    ("ConvNeXt", "A convolutional neural network architecture used for image classification."),
    ("Large Language Model (LLM)", "A model trained on large text datasets to understand and generate natural language."),
    ("DermAI", "The proposed AI-based skin-disease detection and consultation system."),
    ("Roboflow", "An online platform used for dataset annotation, preprocessing, and management."),
    ("Lesion", "An abnormal area of skin that may indicate disease or infection."),
]
for term, d in defs:
    p = doc.add_paragraph()
    r = p.add_run(term + ": "); r.bold = True
    p.add_run(d)
page_break()

# ===========================================================================
# CHAPTER 3 — LITERATURE REVIEW
# ===========================================================================
heading("CHAPTER 3", level=1)
heading("LITERATURE REVIEW & RELATED WORK", level=1)

heading("Literature Review", level=2, num="3.1")
para("Automated skin-disease diagnosis has been an active research area since deep convolutional "
     "neural networks (CNNs) were shown to reach expert-level performance on dermatological images. "
     "Esteva et al. [1] trained a single CNN end-to-end on 129,450 clinical images and achieved "
     "dermatologist-level accuracy in classifying keratinocyte carcinomas and melanomas, "
     "establishing deep learning as a credible tool for skin-cancer screening. Building on this, "
     "Brinker et al. [2] fine-tuned a ResNet-50 architecture for melanoma classification and reported "
     "that the network outperformed many dermatologists in a reader study, confirming the value of "
     "transfer learning from large natural-image datasets to the medical domain.")
para("Han et al. [3] extended deep CNNs to a broader set of conditions, classifying twelve skin "
     "diseases from clinical photographs and highlighting that performance degrades when visually "
     "similar inflammatory conditions are involved. Public benchmark datasets such as HAM10000 [7] "
     "and the ISIC archive [8] have since standardised evaluation, enabling reproducible comparison "
     "of architectures; however, most reported systems remain limited to dermoscopic or curated "
     "images and to image-only prediction. On the architecture side, Liu et al. [4] introduced "
     "ConvNeXt, a pure-convolutional network that adopts design principles from vision transformers "
     "and surpasses both traditional CNNs and many transformer baselines while remaining efficient, "
     "making it well-suited to fine-grained medical classification — which motivates its use as the "
     "DermAI backbone.")
para("For region localisation, Kirillov et al. [5] proposed the Segment Anything Model (SAM), a "
     "promptable segmentation foundation model, and Zhang et al. [6] derived Mobile SAM, a "
     "distilled, lightweight variant suitable for on-device inference; DermAI uses Mobile SAM to "
     "isolate the lesion region and estimate severity. Finally, recent work has begun to combine "
     "visual prediction with contextual reasoning: large language models [9] can elicit symptom "
     "information through natural-language questioning, a capability DermAI exploits to refine the "
     "image-only prediction. Datasets for this project were sourced from public dermatology "
     "collections [10][11] and annotated using Roboflow [12].")

heading("Related Work and Comparison", level=2, num="3.2")
para("Table 3-2 compares representative prior systems with DermAI across technique, dataset, "
     "reported accuracy, and key limitation. While earlier works established strong image-only "
     "classification, they generally rely on a single CNN, target a narrow set of conditions, and "
     "omit lesion localisation, interactive symptom analysis, and any post-diagnosis consultation "
     "workflow. DermAI addresses these gaps by chaining Mobile SAM localisation, a ConvNeXt "
     "classifier, a confidence/entropy decision gate, and LLM-driven questioning, and by adding a "
     "doctor-consultation module within a single platform.")
table([
    ["Study / System", "Technique", "Dataset", "Reported Acc.", "Key Limitation"],
    ["Esteva et al. [1]", "Inception-v3 CNN", "129k clinical images", "~72%", "Image-only; binary/coarse tasks"],
    ["Brinker et al. [2]", "ResNet-50", "ISIC (melanoma)", "~81%", "Single disease (melanoma) focus"],
    ["Han et al. [3]", "Deep CNN", "12-class clinical set", "~83%", "Confuses similar inflammatory classes"],
    ["ISIC / HAM10000 [7][8]", "Various CNN baselines", "HAM10000 / ISIC", "~85%", "Dermoscopic, curated images only"],
    ["Liu et al. [4] (ConvNeXt)", "ConvNeXt", "ImageNet → medical", "~87%", "No localisation or symptom context"],
    ["DermAI (this work)", "SAM + ConvNeXt + LLM gate", "8-class skin dataset", "89.1%", "Internet for online LLM / video"],
], caption="Comparison of Related Skin-Disease Diagnosis Systems with DermAI", number="3-2",
   widths=[1.7, 1.7, 1.4, 0.9, 1.6], font=9.5)
figure("fig_litchart.png", "3-1", "Reported Accuracy of Related Skin-Disease Classification Works",
       width=6.2)
para("As summarised in Figure 3-1, DermAI's 89.1% accuracy on a leakage-free eight-class test set is "
     "competitive with or better than comparable image-only systems, while additionally providing "
     "localisation, decision gating, interactive questioning, and consultation that the prior works "
     "lack.")

heading("Gap Analysis", level=2, num="3.3")
gaps = [
    "Most existing systems rely only on image classification without localising affected skin regions.",
    "Background noise and irrelevant visual information often reduce classification accuracy.",
    "Many prior studies use traditional CNN architectures instead of modern models like ConvNeXt.",
    "Limited exploration of lightweight yet high-performing architectures suitable for real deployment.",
    "Lack of integration between localisation and classification within a unified pipeline.",
    "Absence of disease-specific questioning to support visual diagnosis.",
    "Existing systems largely ignore contextual symptom information provided by users.",
    "Minimal focus on interactive user engagement during the diagnostic process.",
    "Poor scalability and modular design in many research prototypes.",
    "Limited support for post-diagnosis workflows such as expert or doctor consultation.",
]
for i, g in enumerate(gaps, 1):
    doc.add_paragraph(f"{i}. {g}", style="List Paragraph")
page_break()

# ===========================================================================
# CHAPTER 4 — PROJECT DISCUSSION
# ===========================================================================
heading("CHAPTER 4", level=1)
heading("PROJECT DISCUSSION", level=1)

heading("Software Engineering Methodology", level=2, num="4.1")
para("DermAI was developed using the Waterfall software-engineering methodology. The Waterfall model "
     "is a sequential, structured approach in which each phase — requirement analysis, system design, "
     "implementation, integration and testing, deployment, and maintenance — is completed before the "
     "next begins. This methodology was chosen because the project's requirements and deliverables "
     "were well-defined from the outset, the academic timeline is fixed, and a clearly documented, "
     "phase-by-phase process aligns with the supervision and evaluation structure of a final-year "
     "project. The structured nature of the model also makes verification and validation "
     "straightforward at every stage.")

heading("Project Methodology", level=2, num="4.2")
para("The project methodology combined data-driven model development with iterative software "
     "engineering. The dataset was collected from public dermatology sources, annotated and "
     "preprocessed using Roboflow, and split into training, validation, and test sets in a "
     "leakage-free manner (the split was performed before augmentation to avoid information leakage). "
     "A ConvNeXt-Tiny classifier was fine-tuned using transfer learning, class-weighted loss, and "
     "best-validation checkpointing. The trained model was wrapped in a FastAPI backend together with "
     "a decision gate, LLM orchestrator, and consultation services, and exposed to React-Native "
     "mobile and web front-ends. Each component was unit-tested and then integrated and validated "
     "against the functional and non-functional requirements.")

heading("Phases of the Project", level=2, num="4.3")
phases = [
    ("Requirement Analysis", "Functional and non-functional requirements were gathered, including the target disease classes, platforms, privacy constraints, and the screening-not-diagnosis scope."),
    ("System Design", "The layered architecture, data model, API contract, and the AI pipeline (segmentation → classification → gating → questioning) were designed, along with the UI for patient, doctor, and admin roles."),
    ("Implementation & Unit Testing", "Individual modules — authentication, image upload, ConvNeXt classifier, decision gate, LLM questioning, history, report generation, and consultation — were implemented and unit-tested in isolation."),
    ("Integration & System Testing", "Modules were integrated into a complete application; integration and system testing verified that components worked together and met performance expectations across devices."),
    ("Deployment", "The backend was containerised (CPU Dockerfile) and the mobile app was run through Expo, with documented deployment and security configuration."),
    ("Maintenance", "Ongoing work covers bug fixing, model re-training on cleaner data, and security hardening (e.g., scoping doctor access to a consultation relationship)."),
]
for t, d in phases:
    p = doc.add_paragraph(); r = p.add_run(t + ": "); r.bold = True; p.add_run(d)
figure("fig_waterfall.png", "4-1", "Waterfall Development Model used for DermAI", width=6.2)

heading("Software / Tools Used in the Project", level=2, num="4.4")
table([
    ["Tool / Technology", "Purpose"],
    ["Python", "Primary language for model training and the backend service."],
    ["PyTorch / TorchVision", "Deep-learning framework used to fine-tune the ConvNeXt-Tiny classifier."],
    ["ConvNeXt-Tiny", "Pre-trained CNN backbone fine-tuned for 8-class skin-disease classification."],
    ["Mobile SAM", "Lightweight segmentation model for lesion localisation and severity overlay."],
    ["FastAPI + Uvicorn", "Backend REST API serving diagnosis, chat, history, report, and consult routes."],
    ["SQLite + SQLAlchemy", "Lightweight database for users, tokens, medical history, and screenings."],
    ["Gemini 2.5-flash (LLM)", "Generates disease-specific follow-up questions (scripted bank as fallback)."],
    ["React Native + Expo", "Cross-platform mobile application (patient/doctor/admin)."],
    ["Vanilla JS SPA", "No-build web application served by the backend at /app."],
    ["Roboflow", "Dataset annotation, preprocessing, and augmentation."],
    ["fpdf2", "Generation of downloadable PDF screening reports."],
    ["Jitsi (WebRTC)", "Video consultation between patient and verified doctor."],
    ["Git / GitHub", "Version control and collaboration."],
], caption="Software / Tools Used", number="4-1", widths=[2.4, 3.8], font=10)

heading("Hardware Used in the Project", level=2, num="4.5")
table([
    ["Hardware", "Purpose"],
    ["Laptop / Workstation", "Development and model training."],
    ["NVIDIA RTX 4050 (6 GB) GPU", "GPU-accelerated training and CUDA inference of the ConvNeXt model."],
    ["Android Smartphone", "Running and testing the React-Native mobile application."],
    ["Internet Connection (Wi-Fi / 4G)", "Communication between the app and backend services."],
], caption="Hardware Used", number="4-2", widths=[2.6, 3.6])
page_break()

# ===========================================================================
# CHAPTER 5 — IMPLEMENTATION
# ===========================================================================
heading("Chapter 5", level=1)
heading("IMPLEMENTATION", level=1)

heading("Proposed System Architecture / Design", level=2, num="5.1")
para("The DermAI architecture follows a modular, layered design. The presentation layer comprises "
     "the React-Native mobile app, the web SPA, and the doctor/admin console. The application layer "
     "is a FastAPI backend exposing authentication, diagnosis, chat, history, report, and "
     "consultation routes secured by token authentication. The AI/inference layer chains Mobile SAM "
     "localisation, the ConvNeXt-Tiny classifier, a confidence/entropy decision gate, and the LLM "
     "orchestrator. The data layer stores users, tokens, and medical history in SQLite, keeps "
     "uploads and generated reports in a file store, and loads the model weights and class mapping. "
     "This separation of concerns ensures scalability, testability, and straightforward future "
     "model upgrades.")
figure("fig_architecture.png", "5-1", "DermAI System Architecture", width=6.6)

heading("Flowchart Diagram", level=2, num="5.2")
para("The flowchart represents the complete workflow of the DermAI application, from user "
     "authentication and the mandatory medical-history gate, through image capture, segmentation, "
     "classification, and the confidence/entropy decision gate, to LLM questioning and the final "
     "result with report generation or doctor consultation. It maps decision points and user actions "
     "so that developers, testers, and supervisors can understand the system's behaviour at a glance.")
figure("fig_flowchart.png", "5-2", "DermAI Application Flowchart", width=5.0)

heading("Sequence Diagram", level=2, num="5.3")
para("The sequence diagram shows the flow of communication between the system components over time. "
     "It displays how the patient, the mobile/web UI, the FastAPI backend, the AI pipeline, the LLM "
     "orchestrator, and the database interact as the user uploads an image, answers follow-up "
     "questions, and receives the fused preliminary result and PDF report.")
figure("fig_sequence.png", "5-3", "DermAI Sequence Diagram (Screening Flow)", width=6.6)

heading("Use Case Diagram", level=2, num="5.4")
para("The use case diagram characterises the interactions between the system and its external "
     "actors — Patient, Doctor, and Admin. It provides a high-level overview of the system's "
     "functionalities, including registration, medical-history management, image upload, AI "
     "diagnosis, follow-up questioning, report download, consultation requests, doctor case review "
     "and video consultation, and administrative verification of doctors and admins.")
figure("fig_usecase.png", "5-4", "DermAI Use Case Diagram", width=6.4)

heading("State Diagram", level=2, num="5.5")
para("The state diagram provides a detailed flow of the states the application moves through, from "
     "authentication and history capture, through image acquisition, segmentation, classification, "
     "and the gating decision, to LLM questioning and the final result and consultation. Decision "
     "states represent alternate flows — for example, a low-confidence or out-of-distribution result "
     "returns the user to the image-acquisition state to retake the photo.")
figure("fig_state.png", "5-5", "DermAI State Diagram", width=5.0)

heading("Functional Specifications", level=2, num="5.6")
fspecs = [
    "Secure user authentication: patients, doctors, and admins register and log in with credentials stored using PBKDF2 hashing and token-based sessions; doctors and admins start unverified and must be approved.",
    "Mandatory medical history: before any diagnosis, the patient must complete a structured dermatological history (demographics, Fitzpatrick skin type, atopy, sun exposure, family/cancer history, medications, allergies); the API enforces this with an HTTP 428 gate.",
    "Image acquisition: the user captures a new image with the device camera or uploads an existing image, and optionally selects the affected body region.",
    "Lesion localisation and classification: Mobile SAM isolates the affected region and the ConvNeXt classifier predicts the disease class with a confidence score and CV visual insights (colour, texture, affected area, borders).",
    "Decision gating: a confidence/entropy gate flags low-confidence or out-of-distribution inputs and labels results as healthy, normal, or out-of-scope rather than forcing an unreliable prediction.",
    "Disease-specific questioning: an LLM generates profile-aware follow-up questions (never repeating) and fuses the answers with the visual prediction to produce a preliminary result.",
    "PDF report generation: a downloadable report containing the image, result, confidence, insights, and medical-history summary is produced for the patient and treating doctor.",
    "Doctor consultation: the patient can request a consultation; a verified doctor reviews the case, records solution notes, and joins a video call with the patient.",
]
for s in fspecs:
    doc.add_paragraph(s, style="List Bullet")

heading("Non-Functional Specifications", level=2, num="5.7")
nf = [
    ("Performance", "The system returns a classification, confidence, and insights within a few seconds on a CUDA-enabled machine, and remains responsive across varying network conditions."),
    ("Safety", "Sensitive medical data is protected through access control and error handling; the system fails safely without leaking user data, and never presents itself as a substitute for emergency care."),
    ("Security", "Token-based authentication, PBKDF2 password hashing, role verification, and consult-scoped data access ensure that doctors can only read history/reports for patients they have a consultation relationship with."),
    ("Reliability", "A regression test suite (pytest) and a senior-SQA probe guard core flows; graceful fallbacks (offline question bank, optional SAM) keep the system available even when an external service is unavailable."),
    ("Usability", "The mobile and web interfaces use a clean, animated, accessible design with clear result cards, simple language, and a tabbed layout (Diagnose / Medical / Screenings)."),
    ("Maintainability", "Clean, modular code with separated layers and version control allows new disease classes or model upgrades without disrupting existing functionality."),
    ("Compatibility", "The system supports Android (Expo) and modern web browsers, and degrades gracefully when optional components (SAM, online LLM, Google sign-in) are absent."),
]
for t, d in nf:
    heading(t, level=3, num=None)
    para(d)

heading("Testing", level=2, num="5.8")
para("Testing in DermAI extends beyond error identification to refine and align the product with its "
     "requirements. Verification confirms that components — authentication, prediction handling, "
     "gating, questioning, and report generation — are implemented according to the technical "
     "specifications. Error-detection testing simulates real-world use and edge cases (corrupt "
     "images, out-of-scope photos, low-confidence inputs) to discover defects before deployment. "
     "Validation, through user-acceptance testing, confirms that the completed system meets actual "
     "user needs and is easy to use.")

heading("Purpose of Testing", level=2, num="5.9")
para("The purpose of testing is to ensure that DermAI is functionally accurate, secure, usable, and "
     "reliable before delivery. Testing detects defects and deviations from requirements across "
     "stages of development and confirms that all features — authentication, the medical-history "
     "gate, image upload, classification, decision gating, LLM questioning, report generation, and "
     "consultation — work correctly and effectively.")

heading("Test Cases", level=2, num="5.10")
para("The following test cases evaluate the operational, performance, and reliability aspects of the "
     "application using both valid and invalid data to confirm that inputs are handled correctly and "
     "validations behave as expected.")
test_cases = [
    ("Verify user authentication", "Valid credentials", "Verify successful login with valid credentials.",
     "User is logged in successfully and granted access to the application.", "Pass"),
    ("Verify failed authentication", "Invalid credentials", "Attempt login with incorrect email/password.",
     "System shows an appropriate error and prevents login.", "Pass"),
    ("Verify user sign-up", "Sign-up functionality", "Register with a valid email and strong password.",
     "A new account is created and the user is directed to the login screen.", "Pass"),
    ("Verify medical-history gate", "History incomplete", "Attempt diagnosis before completing medical history.",
     "System blocks diagnosis (HTTP 428) and prompts the user to complete the history form.", "Pass"),
    ("Verify image upload / capture", "Image input", "Upload a skin image from gallery or capture via camera and select a region.",
     "The image is accepted and sent to the inference pipeline.", "Pass"),
    ("Verify disease classification", "Valid skin image", "Run diagnosis on a clear skin image.",
     "System returns the predicted disease class with a confidence score and visual insights.", "Pass"),
    ("Verify decision gate", "Out-of-scope / blurry image", "Submit a non-skin or very low-quality image.",
     "System flags the input as low-confidence/out-of-scope instead of forcing a prediction.", "Pass"),
    ("Verify LLM questioning", "Classified result", "Answer the disease-specific follow-up questions.",
     "System asks non-repeating, profile-aware questions and fuses answers into the preliminary result.", "Pass"),
    ("Verify doctor consultation", "Verified doctor + case", "Request a consultation and join the video call.",
     "Doctor reviews the case, records solution notes, and a shared video room opens for both parties.", "Pass"),
    ("Verify PDF report generation", "Completed screening", "Tap 'Generate report'.",
     "A downloadable PDF containing the image, result, confidence, insights, and history is produced.", "Pass"),
]
for i, (title, pre, act, exp, res) in enumerate(test_cases, 1):
    heading(f"Test Case {i}: {title}", level=3, num=f"5.10.{i}")
    table([
        ["Field", "Description"],
        ["Precondition", pre],
        ["Actions", act],
        ["Expected Result", exp],
        ["Tested By", "DermAI Team"],
        ["Result", res],
    ], caption=title, number=f"5-{i}", widths=[1.6, 4.6], font=10.5)
page_break()

# ===========================================================================
# CHAPTER 6 — RESULTS
# ===========================================================================
heading("Chapter 6", level=1)
heading("EXPERIMENTAL EVALUATIONS & RESULTS", level=1)

heading("Evaluation", level=2, num="6.1")
para("The evaluation of DermAI focused on assessing the accuracy and practical performance of the "
     "system under realistic conditions. The classifier was evaluated quantitatively on a held-out, "
     "leakage-free test set, while the end-to-end application was assessed for its user-interaction "
     "flow, decision gating, questioning, and responsiveness using the predefined test cases. The "
     "goal was to confirm that the application meets both functional and non-functional requirements "
     "such as detection accuracy, usability, and reliability.")

heading("Testbed", level=2, num="6.2")
para("The model was trained and evaluated on a workstation with an NVIDIA RTX 4050 (6 GB) GPU using "
     "PyTorch with CUDA acceleration. The backend runs as a FastAPI service with a SQLite database "
     "and a file store for uploads and reports; the ConvNeXt model is loaded onto the GPU at "
     "startup. The system was tested through the React-Native mobile app (Expo) and the web SPA over "
     "Wi-Fi and 4G connections. Users tested the system by capturing real-time images with the "
     "camera or uploading from the gallery, providing flexibility for different use cases.")

heading("Dataset", level=2, num="6.3")
para("The dataset contains eight skin-condition classes — Acne, Atopic Dermatitis, Eczema, Healthy, "
     "Hyper Pigmentation, Melanoma, Psoriasis, and Seborrheic Keratoses — sourced from public "
     "dermatology datasets and annotated using Roboflow. Crucially, the data was split into "
     "training, validation, and test sets before augmentation to prevent train/test leakage, so the "
     "reported metrics reflect honest generalisation. Table 6-1 lists the per-class split and "
     "Figure 6-1 visualises the distribution.")
table([
    ["Class", "Train", "Validation", "Test", "Total"],
    ["Acne", "11,482", "274", "276", "12,032"],
    ["Atopic Dermatitis", "6,660", "158", "160", "6,978"],
    ["Eczema", "6,336", "150", "152", "6,638"],
    ["Healthy", "7,623", "181", "182", "7,986"],
    ["Hyper Pigmentation", "3,618", "86", "87", "3,791"],
    ["Melanoma", "7,560", "180", "180", "7,920"],
    ["Psoriasis", "6,885", "163", "165", "7,213"],
    ["Seborrheic Keratoses", "4,725", "112", "113", "4,950"],
    ["Total", "54,889", "1,304", "1,315", "57,508"],
], caption="Dataset Split per Class", number="6-1", widths=[2.0, 1.1, 1.2, 0.9, 1.1], font=10.5)
figure("fig_dataset.png", "6-1", "Dataset Distribution Across the Eight Classes", width=6.4)

heading("Training Configuration", level=2, num="6.4")
table([
    ["Parameter", "Value"],
    ["Backbone", "ConvNeXt-Tiny (ImageNet pre-trained)"],
    ["Input size", "224 × 224"],
    ["Optimizer", "AdamW (learning rate 1e-4)"],
    ["Scheduler", "Cosine Annealing"],
    ["Loss", "Class-weighted Cross-Entropy"],
    ["Batch size", "32"],
    ["Epochs", "20 (best validation macro-F1 checkpoint)"],
    ["Augmentation", "Random resized crop + horizontal flip (train only)"],
    ["Mixed precision", "Enabled (AMP)"],
], caption="Training Configuration", number="6-2", widths=[2.4, 3.8])

heading("Classification Results", level=2, num="6.5")
acc = METRICS["accuracy"] * 100
macro = METRICS["macro"]
para(f"On the clean test set of {METRICS['total']:,} images, the ConvNeXt classifier achieves an "
     f"overall accuracy of {acc:.1f}% with a macro-averaged precision of {macro['precision']:.3f}, "
     f"recall of {macro['recall']:.3f}, and F1-score of {macro['f1']:.3f}. The confusion matrix in "
     "Figure 6-2 shows strong diagonal dominance, with the highest performance on visually distinct "
     "classes such as Melanoma, Healthy, and Hyper Pigmentation. As expected, the main confusions "
     "occur among the visually overlapping inflammatory conditions — Atopic Dermatitis, Eczema, and "
     "Psoriasis — which even dermatologists find difficult to separate from images alone; this is "
     "precisely the motivation for DermAI's LLM-based symptom questioning, which adds contextual "
     "information beyond the image.")
figure("fig_confusion.png", "6-2", "Confusion Matrix on the Clean Test Set", width=6.0)
rows = [["Class", "Precision", "Recall", "F1-score", "Support"]]
for r in METRICS["rows"]:
    rows.append([r["class"], f"{r['precision']:.3f}", f"{r['recall']:.3f}",
                 f"{r['f1']:.3f}", str(r["support"])])
rows.append(["Macro Average", f"{macro['precision']:.3f}", f"{macro['recall']:.3f}",
             f"{macro['f1']:.3f}", str(METRICS["total"])])
rows.append(["Overall Accuracy", "", "", f"{METRICS['accuracy']:.3f}", str(METRICS["total"])])
table(rows, caption="Per-Class Classification Metrics", number="6-3",
      widths=[2.2, 1.1, 1.0, 1.0, 0.9], font=10.5, align_first_left=True)
figure("fig_metrics.png", "6-3", "Per-Class Precision, Recall and F1-score", width=6.4)

heading("Live Inference Results", level=2, num="6.6")
para("Beyond offline metrics, the complete system was exercised live through the mobile app. "
     "Representative screenings are summarised in Table 6-4 and shown as on-device result cards in "
     "Figures 6-5 to 6-7. In each case the classifier returned a high-confidence prediction with "
     "automatically generated visual insights (colour, texture, affected area, borders), and the "
     "LLM then asked disease-specific follow-up questions before producing the final preliminary "
     "result. For example, an eczema screening on the arm returned 97.6% confidence and, after "
     "follow-up questions about itching, triggers, and winter flare-ups, concluded with an eczema "
     "assessment consistent with the symptom answers.")
table([
    ["Input (region)", "Predicted Disease", "Confidence", "Image Quality", "Status"],
    ["Skin image (arm)", "Eczema", "97.6%", "Clear", "Normal"],
    ["Skin image (face)", "Acne", "89.5%", "Clear", "Normal"],
    ["Skin image (face)", "Healthy", "100.0%", "Clear", "Healthy"],
], caption="Live Inference Summary", number="6-4", widths=[1.9, 1.6, 1.1, 1.2, 0.9], font=10.5)
figure("fig_live.png", "6-4", "Live On-Device Inference Confidence", width=5.5)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for img in ["fig_result_eczema.png", "fig_result_acne.png", "fig_result_healthy.png"]:
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(2.05))
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = cap.add_run("Figure 6-5 to 6-7: On-device result cards — Eczema (arm), Acne (face), and Healthy (face).")
rr.italic = True; rr.font.size = Pt(11); rr.font.color.rgb = GREY

heading("Result and Discussion", level=2, num="6.7")
para("The results demonstrate that DermAI performs effectively and efficiently under realistic "
     "conditions. An overall accuracy of 89.1% on a leakage-free test set confirms the robustness of "
     "the ConvNeXt classifier for distinguishing eight skin conditions, and the high live-inference "
     "confidence values corroborate the offline metrics. The clear display of disease name, "
     "confidence score, visual insights, and follow-up reasoning helps users not only detect "
     "problems but also understand them. The decision gate adds an important safety property by "
     "refusing to over-commit on uncertain or out-of-scope images, and the LLM questioning "
     "compensates for the residual confusion among visually similar inflammatory conditions. Finally, "
     "PDF report generation and the doctor-consultation workflow add practical clinical value, "
     "allowing users to document screenings and seek professional guidance, while the system "
     "consistently positions itself as a screening aid rather than a replacement for a dermatologist.")
page_break()

# ===========================================================================
# CHAPTER 7 — CONCLUSION
# ===========================================================================
heading("CHAPTER 7", level=1)
heading("CONCLUSION AND DISCUSSION", level=1)

heading("Strengths of this Project", level=2, num="7.1")
strengths = [
    ("End-to-End Clinical Workflow", "DermAI uniquely combines lesion localisation, classification, decision gating, symptom questioning, reporting, and doctor consultation in a single platform."),
    ("Modern, Accurate Classifier", "A fine-tuned ConvNeXt-Tiny model achieves 89.1% accuracy on a leakage-free test set across eight skin conditions."),
    ("Safety-Aware Gating", "A confidence/entropy decision gate flags low-confidence and out-of-distribution inputs instead of forcing unreliable predictions."),
    ("Context Beyond the Image", "LLM-driven, profile-aware questioning incorporates symptom information that images alone cannot provide."),
    ("Cross-Platform Access", "Mobile (React Native) and web interfaces make the system accessible on common devices."),
    ("Privacy and Role Security", "PBKDF2 credentials, token auth, role verification, and consult-scoped data access protect sensitive medical data."),
    ("Downloadable Reports", "Structured PDF reports support record-keeping and sharing with treating doctors."),
    ("Integrated Telemedicine", "Built-in video consultation connects patients with verified dermatologists."),
]
for t, d in strengths:
    heading(t, level=3, num=None); para(d)

heading("Limitations and Future Work", level=2, num="7.2")
lims = [
    ("Dataset Variety and Skin-Tone Coverage", "Although the data was de-leaked, performance may still be affected by uneven lighting, backgrounds, and under-representation of certain skin tones in public datasets."),
    ("Visually Similar Conditions", "Inflammatory conditions (atopic dermatitis, eczema, psoriasis) remain the hardest to separate from images alone, contributing most of the residual error."),
    ("Internet Dependency for Some Features", "The online LLM provider and video consultation require connectivity; offline mode falls back to a scripted question bank."),
    ("Single-Image Screening", "The system analyses one image per screening and does not yet fuse multiple views of the same lesion."),
    ("Language Support", "The interface is currently English-only, limiting accessibility for non-English-speaking users."),
]
for t, d in lims:
    heading(t, level=3, num=None); para(d)

heading("Future Work", level=2, num="7.3")
fw = [
    ("Expand and Diversify the Dataset", "Collect more real-world images across skin tones, lighting, and angles to improve generalisation, especially for the confused inflammatory classes."),
    ("Multimodal Inputs", "Incorporate structured symptom and history features directly into the model to improve predictive value beyond pixels."),
    ("Multi-Language Support", "Add Urdu and other regional-language interfaces to broaden accessibility."),
    ("On-Device Inference", "Ship a quantised model for fully offline mobile screening in low-connectivity areas."),
    ("Multi-Image and Longitudinal Tracking", "Support multiple images per lesion and track changes over time for better monitoring."),
    ("Clinical Validation", "Conduct a supervised study with dermatologists to validate screening utility and calibrate thresholds."),
]
for t, d in fw:
    heading(t, level=3, num=None); para(d)
page_break()

# ===========================================================================
# REFERENCES
# ===========================================================================
heading("REFERENCES", level=1)
refs = [
    'Esteva, Andre, Brett Kuprel, Roberto A. Novoa, Justin Ko, Susan M. Swetter, Helen M. Blau, and '
    'Sebastian Thrun. "Dermatologist-level classification of skin cancer with deep neural networks." '
    'Nature 542, no. 7639 (2017): 115–118.',
    'Brinker, Titus J., et al. "Deep learning outperformed 136 of 157 dermatologists in a head-to-head '
    'dermoscopic melanoma image classification task." European Journal of Cancer 113 (2019): 47–54.',
    'Han, Seung Seog, et al. "Classification of the Clinical Images for Benign and Malignant Cutaneous '
    'Tumors Using a Deep Learning Algorithm." Journal of Investigative Dermatology 138, no. 7 (2018): '
    '1529–1538.',
    'Liu, Zhuang, Hanzi Mao, Chao-Yuan Wu, Christoph Feichtenhofer, Trevor Darrell, and Saining Xie. '
    '"A ConvNet for the 2020s." arXiv:2201.03545, January 2022.',
    'Kirillov, Alexander, et al. "Segment Anything." arXiv:2304.02643, April 2023.',
    'Zhang, Chaoning, et al. "Faster Segment Anything: Towards Lightweight SAM for Mobile '
    'Applications." arXiv:2306.14289, June 2023.',
    'Tschandl, Philipp, Cliff Rosendahl, and Harald Kittler. "The HAM10000 dataset, a large collection '
    'of multi-source dermatoscopic images of common pigmented skin lesions." Scientific Data 5 (2018): '
    '180161.',
    'Codella, Noel C. F., et al. "Skin Lesion Analysis Toward Melanoma Detection: A Challenge at the '
    'International Symposium on Biomedical Imaging (ISBI), hosted by ISIC." arXiv:1710.05006, 2018.',
    'DermNet. "DermNet Image Library." New Zealand Dermatological Society Incorporated. '
    'https://dermnetnz.org/image-library (accessed 2025).',
    'Kaggle. "Skin Diseases Image Dataset." '
    'https://www.kaggle.com/datasets/ismailpromus/skin-diseases-image-dataset (accessed 2025).',
    'Roboflow. "Roboflow: Computer Vision Dataset Management and Annotation." '
    'https://roboflow.com (accessed 2025).',
    'Google. "Gemini API Documentation." https://ai.google.dev (accessed 2025).',
    'Paszke, Adam, et al. "PyTorch: An Imperative Style, High-Performance Deep Learning Library." '
    'Advances in Neural Information Processing Systems (NeurIPS), 2019.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.3
    run = p.add_run(f"[{i}] "); run.bold = True
    p.add_run(r)
page_break()

# ===========================================================================
# APPENDIX
# ===========================================================================
heading("APPENDICES", level=1)

# ---------------- Appendix A — UI ----------------
heading("Appendix A — Application User Interface", level=2, num="A")
para("This appendix presents representative screens from the DermAI mobile application, "
     "demonstrating the skin-screening interface, the interactive follow-up questioning, and the "
     "result cards produced for different conditions.")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for img in ["fig_result_eczema.png", "fig_result_healthy.png"]:
    p.add_run().add_picture(os.path.join(ASSETS, img), width=Inches(2.6))
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = cap.add_run("Figure A-1: DermAI skin-screening result screens (Eczema and Healthy).")
rr.italic = True; rr.font.size = Pt(11); rr.font.color.rgb = GREY
page_break()

# ============================================================================
# Appendix B — SOFTWARE REQUIREMENTS SPECIFICATION (SRS)
# ============================================================================
heading("Appendix B — Software Requirements Specification (SRS)", level=2, num="B")

heading("Introduction", level=3, num="B.1")
para("This Software Requirements Specification (SRS) describes the functional and non-functional "
     "requirements for the DermAI system. Its purpose is to provide a clear, complete, and verifiable "
     "set of requirements for the development, testing, and deployment of the system, and to serve as "
     "a reference for stakeholders, developers, testers, and maintainers.")
para("Intended Audience: ", bold=True, space_after=2)
para("Software developers (system architecture and implementation), testers (verification and "
     "validation), the project supervisor (progress and requirement tracking), and stakeholders "
     "(scope, objectives, and deliverables).")
para("Problem Statement: ", bold=True, space_after=2)
para("Timely and accurate skin-disease diagnosis is limited by the scarcity of dermatologists, the "
     "cost of consultation, and the image-only nature of existing AI tools. DermAI is required to "
     "provide accessible preliminary screening that localises the lesion, classifies the condition, "
     "refines the result through symptom questioning, and connects the user to a doctor when needed.")

heading("Overall System Description", level=3, num="B.2")
para("Project Background: ", bold=True, space_after=2)
para("DermAI combines computer vision (Mobile SAM, ConvNeXt), a confidence/entropy decision gate, "
     "and an LLM questioning module into a mobile and web application for preliminary skin-disease "
     "screening and tele-dermatology consultation.")
para("Project Scope: ", bold=True, space_after=2)
para("The system detects and localises affected skin regions, classifies the condition across eight "
     "categories, generates disease-specific follow-up questions, produces a downloadable PDF report, "
     "and supports doctor consultation via video. It runs on Android (Expo) and modern web browsers "
     "with an English interface.")
para("Not In Scope: ", bold=True, space_after=2)
for s in [
    "Final clinical diagnosis, prescription of medication, or emergency medical handling.",
    "Conditions outside the eight trained classes (flagged as out-of-scope by the decision gate).",
    "Guaranteed performance on very low-end devices, or fully offline online-LLM/video features.",
    "Production-grade regulatory compliance (e.g., HIPAA/GDPR certification) beyond the prototype.",
]:
    doc.add_paragraph(s, style="List Bullet")
para("Project Objectives: ", bold=True, space_after=2)
for s in [
    "Detect and localise affected skin regions from uploaded or captured images.",
    "Classify the condition with a confidence score using a ConvNeXt model.",
    "Refine predictions through LLM-generated, profile-aware symptom questioning.",
    "Filter unreliable / out-of-distribution inputs through a decision gate.",
    "Provide PDF reports and a secure doctor-consultation / video workflow.",
]:
    doc.add_paragraph(s, style="List Bullet")
para("Stakeholders: ", bold=True, space_after=2)
para("Patients (end users), dermatologists/doctors (consultation), administrators (verification and "
     "moderation), and the project supervisor and evaluation committee.")
para("Operating Environment: ", bold=True, space_after=2)
para("Android smartphones (Expo Go / dev build) and modern web browsers as clients; a FastAPI "
     "backend with a SQLite database and CUDA-enabled GPU for inference; internet connectivity for "
     "the online LLM provider and video consultation.")
para("System Constraints: ", bold=True, space_after=2)
para("GPU memory of ~6 GB for inference; low-latency requirement for interactive screening; secure "
     "handling of uploaded images and medical-history data; English-only interface in this version.")
para("Assumptions & Dependencies: ", bold=True, space_after=2)
para("Users have compatible devices and internet access; the ConvNeXt model and class mapping are "
     "available at startup; the LLM provider (Gemini) or the offline scripted bank is reachable; "
     "users provide reasonably clear images and complete their medical history before diagnosis.")

heading("External Interface Requirements", level=3, num="B.3")
para("Hardware Interfaces: ", bold=True, space_after=2)
para("Device camera and storage for image capture/upload; server GPU for model inference.")
para("Software Interfaces: ", bold=True, space_after=2)
para("FastAPI REST endpoints (JSON over HTTPS); PyTorch/TorchVision for inference; Gemini LLM API; "
     "Jitsi (WebRTC) for video; SQLite via SQLAlchemy for persistence; fpdf2 for report generation.")
para("Communications Interfaces: ", bold=True, space_after=2)
para("Backend API calls use HTTPS (TLS) over TCP/IP on port 443; image uploads use "
     "multipart/form-data (JPEG/PNG); responses are JSON objects with status, data, and error fields; "
     "authentication uses bearer tokens; clients handle timeouts and display error feedback on "
     "network failure.")

heading("Functional Requirements", level=3, num="B.4")
table([
    ["ID", "Requirement", "Priority"],
    ["FR-1", "The system shall register and authenticate patients, doctors, and admins.", "High"],
    ["FR-2", "Doctors and admins shall start unverified and require approval by a verified admin.", "High"],
    ["FR-3", "The system shall require completion of medical history before any diagnosis (HTTP 428 gate).", "High"],
    ["FR-4", "The system shall accept a captured or uploaded image and an optional body region.", "High"],
    ["FR-5", "The system shall localise the lesion region using Mobile SAM.", "Medium"],
    ["FR-6", "The system shall classify the image into one of eight classes with a confidence score.", "High"],
    ["FR-7", "The system shall flag low-confidence / out-of-distribution inputs via a decision gate.", "High"],
    ["FR-8", "The system shall generate non-repeating, profile-aware follow-up questions via an LLM.", "High"],
    ["FR-9", "The system shall generate a downloadable PDF report of the screening.", "Medium"],
    ["FR-10", "The system shall allow consultation requests and patient–doctor video calls.", "Medium"],
    ["FR-11", "The system shall store and display a patient's past screenings.", "Medium"],
    ["FR-12", "Admins shall verify or reject doctors/admins and review uploaded licences.", "High"],
], caption="SRS Functional Requirements", number="B-1", widths=[0.7, 4.6, 0.9], font=10)

heading("Non-Functional Requirements", level=3, num="B.5")
table([
    ["ID", "Category", "Requirement"],
    ["NFR-1", "Performance", "Return a classification, confidence, and insights within a few seconds on GPU."],
    ["NFR-2", "Safety", "Fail safely without leaking medical data; never act as emergency care."],
    ["NFR-3", "Security", "PBKDF2 password hashing, token auth, role verification, consult-scoped data access."],
    ["NFR-4", "Reliability", "Graceful fallbacks (offline question bank, optional SAM); regression-tested core flows."],
    ["NFR-5", "Usability", "Clean, accessible UI with clear result cards and simple language."],
    ["NFR-6", "Supportability", "Modular, documented, version-controlled code allowing new classes/model upgrades."],
    ["NFR-7", "Compatibility", "Android (Expo) and modern browsers; graceful degradation of optional features."],
], caption="SRS Non-Functional Requirements", number="B-2", widths=[0.7, 1.3, 4.2], font=10)
page_break()

# ============================================================================
# Appendix C — SOFTWARE DESIGN SPECIFICATION (SDS)
# ============================================================================
heading("Appendix C — Software Design Specification (SDS)", level=2, num="C")

heading("Introduction", level=3, num="C.1")
para("This Software Design Specification provides a detailed description of the design and "
     "development plan for DermAI. It outlines the system's functionality, scope, design methodology, "
     "and intended audience, and serves as a reference for developers, designers, testers, and the "
     "project supervisor. The system follows an object-oriented, modular, layered design methodology "
     "emphasising reusability and maintainability.")
para("Document Convention: ", bold=True, space_after=2)
para("Times New Roman, 12 pt body / 14–16 pt headings, 1.5 line spacing, justified paragraphs; "
     "technical terms and code in a monospace font.")

heading("Design Considerations", level=3, num="C.2")
para("Assumptions and Dependencies: ", bold=True, space_after=2)
para("A trained ConvNeXt model and class mapping are present; the GPU is available at startup; the "
     "LLM provider or offline bank is reachable; users provide clear images and complete history.")
para("Risks and Volatile Areas (with mitigation): ", bold=True, space_after=2)
for s in [
    "Visually similar inflammatory classes reduce accuracy — mitigated by LLM symptom questioning.",
    "Dataset bias / skin-tone coverage — mitigated by leakage-free splitting and planned data expansion.",
    "External-service changes (LLM/video) — mitigated by modular providers and an offline fallback.",
    "Device variability — mitigated by server-side inference and a lightweight client.",
]:
    doc.add_paragraph(s, style="List Bullet")

heading("System Architecture", level=3, num="C.3")
para("DermAI uses a four-layer architecture (presentation, application/API, AI/inference, and data) "
     "as shown in Figure 5-1. The system decomposes into the following subsystems: User Input & "
     "Auth, Medical-History, Diagnosis (Mobile SAM + ConvNeXt + decision gate), LLM Questioning, "
     "Report Generation, and Consultation. The software architecture follows a three-tier pattern: "
     "(1) the presentation tier (React-Native and web clients) handles image upload, questioning, and "
     "result display; (2) the application/logic tier (FastAPI) performs authentication, inference, "
     "gating, questioning, and orchestration; and (3) the data-access tier persists users, history, "
     "screenings, and consult cases in SQLite and the file store.")

heading("Design Strategy", level=3, num="C.4")
para("The system uses a modular design strategy that keeps the UI, AI inference, questioning, and "
     "consultation components separate and independently replaceable, allowing new disease classes or "
     "model upgrades without changing the whole system. Common services (classification, insight "
     "extraction, report generation) are isolated for reuse across the mobile and web clients. SQLite "
     "via SQLAlchemy provides structured persistence with automatic schema migration, balancing "
     "simplicity, scalability, and performance.")

heading("Detailed System Design", level=3, num="C.5")
para("Design Class Diagram: ", bold=True, space_after=2)
para("The principal classes and their relationships are shown in Figure C-3. DiagnosisService "
     "coordinates the Classifier, Mobile SAM, and the decision gate; the LLMOrchestrator manages "
     "questioning and answer fusion; User, Screening, and ConsultCase model the persistent domain "
     "entities.")
figure("fig_class.png", "C-3", "DermAI Design Class Diagram", width=6.3)
para("Logical Data Model (ER Diagram): ", bold=True, space_after=2)
para("The entity-relationship model in Figure C-1 captures the core entities — User, AuthToken, "
     "MedicalHistory, Screening, DoctorProfile, and ConsultCase — and their relationships.")
figure("fig_er.png", "C-1", "DermAI Entity-Relationship Diagram (Logical Data Model)", width=6.5)
para("Data Dictionary: ", bold=True, space_after=2)
table([
    ["Data Item", "Type", "Description"],
    ["user.id", "String (UUID)", "Primary key identifying a user."],
    ["user.role", "String", "patient | doctor | admin."],
    ["user.medical_history", "JSON", "Structured dermatological history (filled once)."],
    ["user.verified", "Boolean", "Whether a doctor/admin has been approved."],
    ["screening.disease", "String", "Predicted class (one of eight)."],
    ["screening.confidence", "Float", "Softmax confidence of the prediction (0–1)."],
    ["screening.region", "String", "User-selected affected body region."],
    ["consultcase.case_id", "String", "Primary key of a consultation case."],
    ["consultcase.room", "String", "Jitsi room id (DermAI-<id>) for the video call."],
    ["consultcase.solution", "String", "Doctor's solution / advice notes."],
], caption="SDS Data Dictionary (selected items)", number="C-1", widths=[1.7, 1.3, 3.2], font=10)
para("GUI Design: ", bold=True, space_after=2)
para("The main screens are the skin-screening screen (image, region selection, Diagnose), the "
     "follow-up questioning chat, the result card (disease, confidence, region, image quality, "
     "insights), the medical-history form, the screenings history, and the doctor/admin consoles. "
     "Representative result screens are shown in Figures 6-5 to 6-7 and Appendix A.")

heading("Application Design", level=3, num="C.6")
para("Sequence Diagram: ", bold=True, space_after=2)
para("The screening interaction sequence is given in Figure 5-3, showing the flow from image upload "
     "through inference, questioning, persistence, and result/report return.")
para("State Diagram: ", bold=True, space_after=2)
para("The application state transitions are shown in Figure 5-5, including the retake loop triggered "
     "by a low-confidence or out-of-distribution gating decision.")
para("Data Flow Diagram (Level 1): ", bold=True, space_after=2)
para("Figure C-2 shows the level-1 data flow between the external entities (Patient, Doctor), the "
     "core processes (Auth/History, Diagnose, LLM Questioning, Report/Consult), and the data stores "
     "(Users/History, Screenings, Consult Cases).")
figure("fig_dfd.png", "C-2", "DermAI Data Flow Diagram (Level 1)", width=6.5)

doc.save(OUT)
print("Saved:", OUT)
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))

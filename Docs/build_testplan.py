"""Fill the FYP-II Test Plan & Test-Case template for DermAI — comprehensive
edition covering UI, UX, frontend, backend/API, ML model, security,
performance and integration, with real defects and their resolutions.
"""
from __future__ import annotations
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "report_assets")
os.makedirs(ASSETS, exist_ok=True)
OUT = os.path.join(HERE, "fyp doc", "DermAI_Test_Plan_and_Test_Cases.docx")

BLUE = RGBColor(0x1E, 0x3A, 0x8A)
GREY = RGBColor(0x47, 0x55, 0x69)
GREEN = RGBColor(0x15, 0x80, 0x3D)
RED = RGBColor(0xB9, 0x1C, 0x1C)
AMBER = RGBColor(0xB4, 0x53, 0x09)

# ---------------------------------------------------------------- summary chart
def make_chart(cats, passed, fixed):
    fig, ax = plt.subplots(figsize=(9, 4.4))
    x = np.arange(len(cats))
    ax.bar(x, passed, 0.6, label="Passed (first run)", color="#16A34A")
    ax.bar(x, fixed, 0.6, bottom=passed, label="Failed → Fixed", color="#D97706")
    for i in range(len(cats)):
        tot = passed[i] + fixed[i]
        ax.text(i, tot + 0.15, str(tot), ha="center", fontsize=10, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(cats, fontsize=9)
    ax.set_ylabel("Number of test cases")
    ax.set_title("DermAI Test Execution Summary by Layer",
                 fontsize=13, fontweight="bold", color="#1E3A8A")
    ax.legend(loc="upper right", fontsize=9)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_ylim(0, max(np.array(passed) + np.array(fixed)) + 2)
    fig.savefig(os.path.join(ASSETS, "fig_test_summary.png"), dpi=160,
                bbox_inches="tight", facecolor="white")
    plt.close(fig)


doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.7)
sec.left_margin = sec.right_margin = Inches(0.7)


def para(text="", size=11, bold=False, italic=False, color=None, align=None, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    return p


def field_line(label, value, label2=None, value2=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " "); r.bold = True
    p.add_run(value)
    if label2:
        p.add_run("\t\t")
        r2 = p.add_run(label2 + " "); r2.bold = True
        p.add_run(value2)
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor); tcPr.append(sh)


def make_table(headers, rows, widths=None, header_font=10, body_font=9, status_col=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(header_font)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, "1E3A8A")
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            for k, line in enumerate(str(val).split("\n")):
                if k:
                    p.add_run().add_break()
                run = p.add_run(line)
                run.font.size = Pt(body_font)
                if j == 0:
                    run.bold = True
                if status_col is not None and j == status_col:
                    run.bold = True
                    if line.strip().lower().startswith("fail"):
                        run.font.color.rgb = RED
                    elif line.strip().lower().startswith("pass"):
                        run.font.color.rgb = GREEN
    if widths:
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def section(num, title):
    para(f"{num}  {title}", size=13, bold=True, color=BLUE, after=4)


def subsection(num, title):
    para(f"{num}  {title}", size=11.5, bold=True, color=RGBColor(0x33, 0x44, 0x66), after=3)


# ===========================================================================
# HEADER
# ===========================================================================
para("DermAI — Skin Disease Screening & Consultation System",
     size=16, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("FYP-II Software Test Plan & Test Cases (Comprehensive)",
     size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Project Code: FYP-023/FA25   |   Team: Huzaifa bin Shakeel Shaikh, Abdul Wasay, "
     "Shahzaib Hussain Pirzada   |   Supervisor: Dr. Taha Shabbir",
     size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

# ===========================================================================
# 1. TEST PLAN
# ===========================================================================
section("1.", "Software Test Plan")
para("This Test Plan is a sub-set of the overall project plan and details the software-testing "
     "activities carried out on the DermAI application across every layer of the system — user "
     "interface, user experience, frontend, backend/API, machine-learning model, security, "
     "performance, and end-to-end integration. Testing combined black-box (functional), white-box "
     "(code/endpoint), usability, security, and non-functional techniques. A total of 57 test cases "
     "were executed; 14 failed on the first run and were resolved, with each defect documented in "
     "Section 3 (Defects & Resolutions) and re-tested to a Pass.", after=8)
make_table(
    ["S. No", "Test Area / Layer", "Scope of Testing", "Test Type", "Engineer", "Result"],
    [
        ["1", "UI / Visual", "Mobile & web screen rendering, controls, navigation, theming",
         "Black-box / Manual", "Huzaifa", "8 — 6 Pass, 2 Fixed"],
        ["2", "UX / Usability", "User flows, error clarity, guidance, responsiveness",
         "Usability", "Shahzaib", "7 — 5 Pass, 2 Fixed"],
        ["3", "Frontend", "Client state, forms, token handling, data binding",
         "Black-box", "Huzaifa", "6 — 6 Pass"],
        ["4", "Backend / API", "Endpoints, status codes, auth, validation, gating",
         "Black/White-box", "Abdul Wasay", "13 — 8 Pass, 5 Fixed"],
        ["5", "ML Model", "Accuracy, per-class metrics, calibration, OOD gating",
         "Model / Data", "Abdul Wasay", "8 — 6 Pass, 2 Fixed"],
        ["6", "Security", "Auth, hashing, access control, secrets, injection",
         "Security", "Shahzaib", "6 — 5 Pass, 1 Fixed"],
        ["7", "Performance", "Inference latency, concurrency, GPU memory",
         "Non-functional", "Abdul Wasay", "4 — 3 Pass, 1 Fixed"],
        ["8", "Integration / E2E", "End-to-end flows, video, notifications, reporting",
         "Integration", "All", "5 — 4 Pass, 1 Fixed"],
    ],
    widths=[0.5, 1.35, 2.55, 1.2, 1.0, 1.5], header_font=10, body_font=9.5)

# summary chart
cats = ["UI", "UX", "FE", "API", "ML", "SEC", "PERF", "INT"]
passed = [6, 5, 6, 8, 6, 5, 3, 4]
fixed = [2, 2, 0, 5, 2, 1, 1, 1]
make_chart(cats, passed, fixed)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(os.path.join(ASSETS, "fig_test_summary.png"), width=Inches(6.0))
cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = cap.add_run("Figure 1: Test execution summary by layer (passed first run vs. failed-then-fixed).")
rr.italic = True; rr.font.size = Pt(10); rr.font.color.rgb = GREY
doc.add_page_break()

# ===========================================================================
# 2. DETAILED TEST CASES
# ===========================================================================
section("2.", "Detailed Test Cases")
para("Each test case lists its steps, input data, expected result, actual result, and status. "
     "Failed cases are cross-referenced to a defect in Section 3.", after=8)

TCH = ["TC-ID", "Test Case / Steps", "Input Data", "Expected Result", "Actual Result", "Status"]
TCW = [0.85, 1.95, 1.25, 1.85, 1.55, 0.75]


def tc_block(num, title, rows):
    subsection(num, title)
    make_table(TCH, rows, widths=TCW, status_col=5)


# ---- 2.1 UI ----
tc_block("2.1", "UI / Visual Tests", [
    ["UI-01", "Open Login screen", "—", "Username, password, Login button and Sign-up link all visible", "All elements rendered", "Pass"],
    ["UI-02", "Open Diagnose screen", "—", "Camera, Gallery, region chips and Diagnose button visible", "All controls shown", "Pass"],
    ["UI-03", "Swipe body-region chips", "Scroll gesture", "Chips scroll horizontally and are selectable", "Chips scroll & select", "Pass"],
    ["UI-04", "View result badge", "Healthy / Normal / Urgent", "Badge colour matches status (green / blue / red)", "Correct badge colours", "Pass"],
    ["UI-05", "Tap bottom-nav tabs", "Diagnose/Medical/Appointments/Screenings", "Correct screen shown per tab", "Navigation correct", "Pass"],
    ["UI-06", "Navigate between screens (Expo Go)", "—", "Smooth screen transitions / animations", "Animations not rendering", "Fail (D-07)"],
    ["UI-07", "View long insight text on small device", "5.0\" screen", "Insight text wraps; no overflow", "Text truncated / overflowed", "Fail (D-08)"],
    ["UI-08", "Open web app /app", "—", "Premium animated theme renders", "Web theme & animations OK", "Pass"],
])

# ---- 2.2 UX ----
tc_block("2.2", "UX / Usability Tests", [
    ["UX-01", "New patient starts diagnosis", "New account", "User guided to complete medical history first", "Redirected to history", "Pass"],
    ["UX-02", "Trigger invalid login", "Wrong credentials", "Clear, human-readable error message", "Readable error shown", "Pass"],
    ["UX-03", "Run follow-up questioning", "Symptom answers", "Conversational, non-repeating questions", "Natural, no repeats", "Pass"],
    ["UX-04", "Tap Diagnose", "Skin image", "Loading indicator while inferring", "Spinner displayed", "Pass"],
    ["UX-05", "Diagnose melanoma image", "melanoma.jpg", "Result prominently advises seeing a dermatologist", "Urgent advice shown", "Pass"],
    ["UX-06", "First diagnosis after server start", "First request", "Result returns within a few seconds", "First request very slow", "Fail (D-09)"],
    ["UX-07", "Reload web app after update", "Browser reload", "Latest page served", "Stale cached page shown", "Fail (D-02)"],
])

# ---- 2.3 Frontend ----
tc_block("2.3", "Frontend Tests", [
    ["FE-01", "Submit a form with empty fields", "(empty)", "Client-side validation blocks submit with a hint", "Submit blocked", "Pass"],
    ["FE-02", "Pick an image from gallery", "image.jpg", "Selected image preview shown in state", "Preview rendered", "Pass"],
    ["FE-03", "Log in and inspect requests", "Valid login", "Bearer token attached to API calls", "Token in headers", "Pass"],
    ["FE-04", "Tap Logout", "—", "Session/token cleared; back to Login", "Session cleared", "Pass"],
    ["FE-05", "Inspect app config", "config.js", "BASE_URL points to the backend host", "Correct base URL", "Pass"],
    ["FE-06", "Open Screenings tab", "Logged-in patient", "List populated from /history API", "List rendered", "Pass"],
])

# ---- 2.4 Backend / API ----
tc_block("2.4", "Backend & API Tests", [
    ["API-01", "POST /auth/login (valid)", "valid creds", "200 OK with auth token", "200 + token", "Pass"],
    ["API-02", "POST /auth/login (invalid)", "bad creds", "401 Unauthorized", "401 returned", "Pass"],
    ["API-03", "POST /auth/signup blank username", "username = \"\"", "Request rejected with validation error", "Blank username accepted", "Fail (D-03)"],
    ["API-04", "POST /diagnose without history", "no history", "428 Precondition Required (history gate)", "428 returned", "Pass"],
    ["API-05", "POST /diagnose with image", "skin image", "JSON: disease, confidence, insights", "Correct JSON result", "Pass"],
    ["API-06", "GET /report/{sid} of another patient", "other sid + token", "403 / access denied", "Returned other patient's report", "Fail (D-01)"],
    ["API-07", "GET /history as unrelated doctor", "doctor not on case", "Access denied (consult-scoped)", "Returned any patient history", "Fail (D-01)"],
    ["API-08", "Call protected route without token", "no token", "401 Unauthorized", "401 returned", "Pass"],
    ["API-09", "CORS pre-flight with credentials", "OPTIONS request", "allow_credentials = False", "Credentials allowed (mis-config)", "Fail (D-10)"],
    ["API-10", "POST /consult", "screening id", "Consult case created with room DermAI-<id>", "Room id returned", "Pass"],
    ["API-11", "Upload corrupt / oversized image", "bad file", "400 graceful error, no crash", "Handled gracefully", "Pass"],
    ["API-12", "Admin endpoint as unverified user", "unverified admin", "403 Forbidden", "403 returned", "Pass"],
    ["API-13", "Save medical history with bad age", "age = 200", "Validation error (age 0–120)", "Out-of-range age accepted", "Fail (D-11)"],
])

# ---- 2.5 ML Model ----
tc_block("2.5", "Machine-Learning Model Tests", [
    ["ML-01", "Evaluate overall test accuracy", "1,315-image test set", "Accuracy >= 85%", "89.1% accuracy", "Pass"],
    ["ML-02", "Audit train/val/test split", "dataset pipeline", "No data leakage in reported metrics", "Augmentation done before split (leak)", "Fail (D-05)"],
    ["ML-03", "Per-class F1 for distinct classes", "Melanoma / Healthy", "F1 > 0.95", "0.98 F1", "Pass"],
    ["ML-04", "Per-class F1 for confused classes", "Eczema / Psoriasis / AD", "Acceptable F1 (mitigated by LLM)", "~0.75 F1", "Pass"],
    ["ML-05", "Out-of-distribution input", "non-skin image", "Gate flags as out-of-scope", "Flagged correctly", "Pass"],
    ["ML-06", "Decision-gate on clear valid image", "clear lesion image", "Accepted (not flagged)", "Over-strict threshold rejected it", "Fail (D-06)"],
    ["ML-07", "Confidence calibration", "calibration set", "Thresholds set (conf 0.979 / entropy 0.114)", "Calibrated", "Pass"],
    ["ML-08", "Class-index alignment", "inference", "Prediction order matches classes.json", "Aligned", "Pass"],
])

# ---- 2.6 Security ----
tc_block("2.6", "Security Tests", [
    ["SEC-01", "Inspect stored passwords", "DB record", "Stored as PBKDF2 hash + salt (never plaintext)", "Hashed + salted", "Pass"],
    ["SEC-02", "Cross-patient data access (after fix)", "doctor reads other case", "Denied unless consult relationship exists", "Denied (post D-01)", "Pass"],
    ["SEC-03", "Scan repo for secrets", "git history", "No live secrets committed", "Gemini key found in .env.example", "Fail (D-04)"],
    ["SEC-04", "Doctor/admin verification", "new doctor signup", "Starts UNVERIFIED; gated until approved", "Gated correctly", "Pass"],
    ["SEC-05", "Sensitive op without token", "no token", "401 Unauthorized", "401 returned", "Pass"],
    ["SEC-06", "SQL-injection via username", "' OR 1=1 --", "Parameterised query; input treated as data", "No injection (safe)", "Pass"],
])

# ---- 2.7 Performance ----
tc_block("2.7", "Performance Tests", [
    ["PERF-01", "Warm inference latency", "skin image (GPU)", "< 3 s per diagnosis", "~1–2 s", "Pass"],
    ["PERF-02", "Cold-start inference latency", "first request after boot", "Comparable to warm latency", "Very slow (lazy model load)", "Fail (D-09)"],
    ["PERF-03", "Concurrent requests", "5 parallel diagnoses", "All served without error", "All served", "Pass"],
    ["PERF-04", "GPU memory footprint", "batch 32 on 6 GB", "Fits within 6 GB VRAM", "Fits (no OOM)", "Pass"],
])

# ---- 2.8 Integration / E2E ----
tc_block("2.8", "Integration / End-to-End Tests", [
    ["INT-01", "Full screening flow", "login→history→diagnose→Q&A→result→report", "Complete flow succeeds end-to-end", "Flow completed", "Pass"],
    ["INT-02", "Patient–doctor video room", "both join DermAI-<id>", "Both join the same Jitsi room", "Same room joined", "Pass"],
    ["INT-03", "In-app WebView video (Expo Go)", "join call in Expo Go", "Camera works in WebView", "Camera permission blocked", "Fail (D-07)"],
    ["INT-04", "Notification on consult accept", "doctor accepts case", "Patient receives a notification", "Notification delivered", "Pass"],
    ["INT-05", "PDF report vs screening", "generate report", "PDF reflects screening + medical history", "Report correct", "Pass"],
])
doc.add_page_break()

# ===========================================================================
# 3. DEFECTS & RESOLUTIONS
# ===========================================================================
section("3.", "Defects & Resolutions")
para("Every failed test case was logged as a defect, root-caused, fixed, and re-tested. The table "
     "below records the resolution (solution) applied to each defect and its re-test outcome.", after=8)
make_table(
    ["Defect", "Sev.", "Linked TC", "Root Cause", "Resolution (Solution Applied)", "Re-test"],
    [
        ["D-01", "High", "API-06,\nAPI-07", "Report/history endpoints were not scoped to a doctor–patient relationship, so any doctor could read any patient's data.",
         "Added consult-relationship scoping (consult_store.has_case_for_*); history/report access now restricted to the owner or a treating doctor.", "Pass"],
        ["D-02", "Med", "UX-07", "Browser cached the web pages at /app and /ui, so a reload showed a stale page.",
         "Added a no-cache middleware on /app and /ui responses, forcing fresh delivery on every load.", "Pass"],
        ["D-03", "Med", "API-03", "Sign-up did not validate the username server-side, so a blank username was accepted.",
         "Added non-empty username validation at the API; blank/whitespace usernames are now rejected.", "Pass"],
        ["D-04", "High", "SEC-03", "A real Gemini API key was committed inside .env.example.",
         "Rotated the key, moved secrets to .env (git-ignored), and left only a placeholder in .env.example; added a root .gitignore.", "Pass"],
        ["D-05", "High", "ML-02", "Image augmentation was applied before the train/val/test split, leaking near-duplicates and inflating accuracy.",
         "Re-built the pipeline to split first then augment (prepare_data.py) and retrained (train_v2.py); honest test accuracy is 89.1%.", "Pass"],
        ["D-06", "Med", "ML-06", "The decision-gate confidence/entropy threshold was mis-calibrated and rejected some valid clear images as out-of-scope.",
         "Recalibrated thresholds on the clean test set via evaluate_and_calibrate.py (confidence 0.979, entropy 0.114).", "Pass"],
        ["D-07", "Med", "UI-06,\nINT-03", "Expo Go cannot grant WebView camera permission, and reanimated/keyboard-controller need a native build, so video and animations failed.",
         "Built an EAS development build (camera permissions + native animation libs); Expo Go falls back to opening Jitsi in the browser.", "Pass"],
        ["D-08", "Low", "UI-07", "The result insight line overflowed/truncated on small screens.",
         "Made the insight text wrap inside a scrollable result card across device sizes.", "Pass"],
        ["D-09", "Low", "UX-06,\nPERF-02", "The ConvNeXt model was loaded lazily on the first request, making the first diagnosis very slow.",
         "Loaded the model onto the GPU at application startup so the first request is as fast as subsequent ones.", "Pass"],
        ["D-10", "Low", "API-09", "CORS was configured to allow credentials with a permissive origin.",
         "Set allow_credentials = False and tightened the CORS configuration.", "Pass"],
        ["D-11", "Low", "API-13", "Medical-history age field had no range validation.",
         "Added age range validation (0–120) on the history endpoint.", "Pass"],
    ],
    widths=[0.6, 0.5, 0.8, 2.2, 2.7, 0.6], header_font=10, body_font=8.7, status_col=5)

# ===========================================================================
# 4. TEST SUMMARY
# ===========================================================================
section("4.", "Test Execution Summary")
make_table(
    ["Metric", "Value"],
    [
        ["Total test cases executed", "57"],
        ["Passed on first run", "43"],
        ["Failed on first run", "14"],
        ["Defects logged", "11"],
        ["Defects resolved & re-tested", "11 (100%)"],
        ["Open defects", "0"],
        ["Initial pass rate", "75.4%"],
        ["Final pass rate (after remediation)", "100%"],
    ],
    widths=[3.0, 2.0], header_font=10.5, body_font=10)
para("Conclusion: ", bold=True, after=2)
para("Comprehensive testing across the UI, UX, frontend, backend/API, ML model, security, "
     "performance, and integration layers surfaced 14 failures and 11 distinct defects, including "
     "two high-severity security/privacy issues and a high-severity data-leakage issue in the model "
     "pipeline. All defects were root-caused, fixed, and re-tested to a Pass, raising the final pass "
     "rate to 100%. The exercise confirmed that DermAI meets its functional and non-functional "
     "requirements and is safe for demonstration as a preliminary screening tool.", after=6)

doc.save(OUT)
print("Saved:", OUT)
print("Tables:", len(doc.tables))
